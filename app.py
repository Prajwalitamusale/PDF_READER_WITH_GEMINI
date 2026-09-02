import io
import json
import os
import re

import google.generativeai as genai
import streamlit as st
from docx import Document
from pypdf import PdfReader

MAX_FILE_BYTES = 3 * 1024 * 1024
MAX_TEXT_CHARS = 12_000
MAX_ANALYSES_PER_SESSION = 15
MODELS_TO_TRY = (
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
    "gemini-3.5-flash-lite",
    "gemini-flash-latest",
)

CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,500;9..144,600&family=DM+Sans:ital,wght@0,400;0,500;0,600;1,400&display=swap');

html, body, [data-testid="stAppViewContainer"], .stApp {
  background: #0A0C0F !important;
  color: #F3F0E8;
  font-family: "DM Sans", sans-serif;
}
[data-testid="stHeader"] { background: transparent !important; }
[data-testid="stToolbar"], [data-testid="stDecoration"] { display: none !important; }
#MainMenu, footer, [data-testid="stStatusWidget"] { visibility: hidden; }
.block-container { padding-top: 1.4rem; max-width: 980px; }

h1, h2, h3, .hero-title {
  font-family: "Fraunces", Georgia, serif;
  letter-spacing: -0.03em;
}

.nav {
  display: flex; align-items: center; justify-content: space-between;
  padding: 4px 0 18px 0; border-bottom: 1px solid #232833; margin-bottom: 28px;
}
.nav-brand { font-family: "Fraunces", serif; font-size: 1.35rem; font-weight: 600; color: #F3F0E8; }
.nav-meta { font-size: 0.8rem; color: #9AA3B2; }

.hero-kicker {
  display: inline-block; font-size: 0.75rem; font-weight: 600; letter-spacing: 0.08em;
  text-transform: uppercase; color: #0A0C0F; background: #D4F562; padding: 4px 10px; border-radius: 999px;
  margin-bottom: 14px;
}
.hero-title { font-size: 3.15rem !important; line-height: 1.08 !important; margin: 0 0 14px 0; color: #F3F0E8; }
.hero-title em { font-style: italic; color: #D4F562; font-weight: 500; }
.hero-sub { font-size: 1.05rem; color: #B7BEC9; max-width: 640px; line-height: 1.5; margin-bottom: 8px; }

.stats { display: grid; grid-template-columns: repeat(4, 1fr); gap: 10px; margin: 22px 0 8px 0; }
.stat {
  background: #14181F; border: 1px solid #232833; border-radius: 14px; padding: 16px 14px;
}
.stat b { display: block; font-family: "Fraunces", serif; font-size: 1.45rem; color: #D4F562; }
.stat span { font-size: 0.78rem; color: #9AA3B2; }

.section-label {
  font-size: 0.72rem; letter-spacing: 0.12em; text-transform: uppercase; color: #9AA3B2;
  margin: 22px 0 10px 0;
}
.cards { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }
.card {
  background: #14181F; border: 1px solid #232833; border-radius: 16px; padding: 18px 18px 16px 18px;
}
.card h4 { margin: 0 0 6px 0; font-size: 1.05rem; color: #F3F0E8; font-family: "Fraunces", serif; }
.card p { margin: 0; color: #9AA3B2; font-size: 0.9rem; line-height: 1.45; }
.card .tag { font-size: 0.7rem; color: #D4F562; font-weight: 600; margin-bottom: 8px; }

.steps { display: grid; grid-template-columns: repeat(4, 1fr); gap: 10px; }
.step { background: #14181F; border-radius: 16px; padding: 16px; border: 1px solid #232833; }
.step .n { color: #D4F562; font-family: "Fraunces", serif; font-size: 1.2rem; }
.step h4 { margin: 6px 0 4px 0; font-size: 0.95rem; }
.step p { margin: 0; color: #9AA3B2; font-size: 0.8rem; }

.panel {
  background: #14181F; border: 1px solid #232833; border-radius: 20px;
  padding: 8px 8px 18px 8px; margin-top: 8px;
}
.score-wrap {
  background: linear-gradient(180deg, #1A2114 0%, #14181F 70%);
  border: 1px solid #3A4A22; border-radius: 20px; padding: 22px 24px; margin: 8px 0 16px 0;
}
.score-num { font-family: "Fraunces", serif; font-size: 3.4rem; line-height: 1; color: #D4F562; }
.score-label { color: #9AA3B2; font-size: 0.85rem; margin-top: 4px; }
.chips { display: flex; flex-wrap: wrap; gap: 8px; }
.chip {
  border-radius: 999px; padding: 6px 12px; font-size: 0.82rem; border: 1px solid #232833;
}
.chip-in { background: #1A2A1C; color: #C8E6B0; border-color: #2E4A30; }
.chip-out { background: #2A1C16; color: #F0C4B0; border-color: #4A3228; }
.q-item, .s-item {
  background: #0A0C0F; border: 1px solid #232833; border-radius: 12px;
  padding: 12px 14px; margin-bottom: 8px; font-size: 0.92rem; color: #E6E3DA;
}
.q-item b { color: #D4F562; margin-right: 8px; }
.role-pill {
  display: inline-block; margin: 0 8px 8px 0; padding: 8px 14px; border-radius: 999px;
  background: #0A0C0F; border: 1px solid #D4F562; color: #D4F562; font-size: 0.88rem;
}
.foot { color: #6E7684; font-size: 0.75rem; margin-top: 28px; }

div[data-testid="stFileUploader"] section { background: #0A0C0F; border-radius: 12px; }
.stButton > button {
  background: #D4F562 !important; color: #0A0C0F !important; border: 0 !important;
  font-weight: 600 !important; border-radius: 999px !important; padding: 0.5rem 1.4rem !important;
}
.stButton > button:hover { filter: brightness(1.05); }

@media (max-width: 800px) {
  .hero-title { font-size: 2.1rem; }
  .stats, .cards, .steps { grid-template-columns: 1fr 1fr; }
}
@media (max-width: 520px) {
  .stats, .cards, .steps { grid-template-columns: 1fr; }
}
</style>
"""


def inject_css() -> None:
    st.markdown(CSS, unsafe_allow_html=True)


def secret(name: str, default: str = "") -> str:
    try:
        value = st.secrets.get(name)
        if value:
            return str(value)
    except Exception:
        pass
    return os.getenv(name, default)


def extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs)


def extract_resume(uploaded) -> str:
    name = (uploaded.name or "").lower()
    data = uploaded.getvalue()
    if len(data) > MAX_FILE_BYTES:
        raise ValueError("File is larger than 3 MB. Upload a smaller resume.")
    if name.endswith(".pdf"):
        text = extract_pdf(data)
    elif name.endswith(".docx"):
        text = extract_docx(data)
    else:
        raise ValueError("Use a PDF or DOCX resume.")
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) < 80:
        raise ValueError("Could not read enough text from the file. Try another export of the resume.")
    return text[:MAX_TEXT_CHARS]


def parse_json_payload(raw: str) -> dict:
    raw = (raw or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


def analyze(resume_text: str, jd_text: str) -> dict:
    api_key = secret("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_API_KEY is missing. Add it in Streamlit Secrets.")
    genai.configure(api_key=api_key)

    prompt = f"""You are JDFit, an India-first resume vs job-description analyst.
Hiring context: Indian IT hiring (freshers to mid-level), Naukri-style JDs, degrees like B.E./B.Tech/M.Tech, internships, notice period, CTC.
This is an AI estimate, not Naukri, Workday, or any real ATS.

Return ONLY valid JSON with this shape:
{{
  "ats_score": <integer 0-100>,
  "score_summary": "<2 sentences>",
  "matched_skills": ["<skill>", "..."],
  "missing_skills": ["<skill>", "..."],
  "resume_suggestions": ["<actionable edit>", "..."],
  "interview_questions": ["<question>", "..."],
  "recommended_job_roles": ["<role title>", "..."],
  "ats_friendly_resume": "<full plain-text resume, use \\\\n for new lines>"
}}

Rules:
- Score how well THIS resume matches THIS JD (keywords, skills, seniority, tools).
- Prefer skills recruiters search (languages, frameworks, SQL, cloud, DSA) over fluffy adjectives.
- 5-10 matched_skills, 5-10 missing_skills, 5-8 resume_suggestions, 8-12 interview_questions (mix behavioral and technical), 3 recommended_job_roles.
- Suggestions must tell the candidate what to add or rewrite. No generic "be confident".
- If the JD is outside India, still analyze it; keep advice practical.
- ats_friendly_resume: rewrite the SAME person for THIS JD in ATS-safe plain text.
  Use headings exactly: CONTACT, PROFESSIONAL SUMMARY, SKILLS, EXPERIENCE, EDUCATION, PROJECTS (omit a heading if the original has no data).
  Single column. No tables, columns, icons, or photos. Standard bullets with hyphen.
  Mirror JD keywords only where the original resume supports them. Do not invent employers, titles, dates, degrees, or metrics.
  Keep Indian degree names and notice/CTC only if they already appear. Target 400-700 words.

RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text[:MAX_TEXT_CHARS]}
"""

    last_error = None
    for model_name in MODELS_TO_TRY:
        try:
            model = genai.GenerativeModel(
                model_name,
                generation_config={
                    "temperature": 0.3,
                    "response_mime_type": "application/json",
                    "max_output_tokens": 8192,
                },
            )
            response = model.generate_content(prompt)
            payload = parse_json_payload(response.text)
            payload["model_used"] = model_name
            return payload
        except Exception as exc:
            last_error = exc
            continue
    raise RuntimeError(f"Gemini did not return usable JSON. Last error: {last_error}")


def chips(items, kind: str) -> str:
    cls = "chip chip-in" if kind == "in" else "chip chip-out"
    bits = "".join(f'<span class="{cls}">{item}</span>' for item in items or [])
    return f'<div class="chips">{bits}</div>'


def resume_to_docx(text: str) -> bytes:
    doc = Document()
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = line.rstrip(":").upper()
        is_heading = heading in {
            "CONTACT",
            "PROFESSIONAL SUMMARY",
            "SKILLS",
            "EXPERIENCE",
            "EDUCATION",
            "PROJECTS",
        } or (line.isupper() and len(line) < 40)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        if is_heading:
            run.bold = True
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def gate_password() -> bool:
    expected = secret("APP_PASSWORD")
    if not expected:
        return True
    if st.session_state.get("jdfit_ok"):
        return True
    st.markdown(
        """
        <div class="nav"><div class="nav-brand">JDFit</div><div class="nav-meta">Private beta</div></div>
        <div class="hero-kicker">Tester access</div>
        <p class="hero-title">Enter the <em>shared password</em> to continue.</p>
        <p class="hero-sub">This build is invite-only so the free Gemini quota lasts for the first testers.</p>
        """,
        unsafe_allow_html=True,
    )
    entered = st.text_input("Password", type="password", label_visibility="collapsed", placeholder="Password")
    if st.button("Continue"):
        if entered == expected:
            st.session_state["jdfit_ok"] = True
            st.rerun()
        st.error("Wrong password.")
    return False


def render_hero() -> None:
    st.markdown(
        """
        <div class="nav">
          <div class="nav-brand">JDFit</div>
          <div class="nav-meta">Resume to interview-ready · India</div>
        </div>
        <div class="hero-kicker">AI career toolkit · MVP</div>
        <p class="hero-title">Go from resume to<br><em>interview-ready</em> in seconds.</p>
        <p class="hero-sub">
          Upload your resume, paste the job. Get a fit score, missing keywords,
          an ATS-friendly rewrite, and interview questions — for Indian IT hiring.
        </p>
        <div class="stats">
          <div class="stat"><b>1 click</b><span>full analysis</span></div>
          <div class="stat"><b>&lt; 30s</b><span>typical wait</span></div>
          <div class="stat"><b>Free</b><span>to test this MVP</span></div>
          <div class="stat"><b>Invite</b><span>shared tester password</span></div>
        </div>
        <div class="section-label">Try it now</div>
        """,
        unsafe_allow_html=True,
    )


def render_below_fold() -> None:
    st.markdown(
        """
        <div class="section-label">In this version</div>
        <div class="cards">
          <div class="card"><div class="tag">Score</div><h4>Fit audit</h4><p>Original resume vs this JD. See the match score before you apply.</p></div>
          <div class="card"><div class="tag">Keywords</div><h4>Search match</h4><p>Matched vs missing skills side by side — the words recruiters actually search.</p></div>
          <div class="card"><div class="tag">Resume</div><h4>Edit suggestions</h4><p>What to add or rewrite, plus a full ATS-safe version you can download.</p></div>
          <div class="card"><div class="tag">Interview</div><h4>Practice kit</h4><p>Behavioral and technical questions aimed at this role, not a generic bank.</p></div>
          <div class="card"><div class="tag">ATS</div><h4>ATS-friendly resume</h4><p>Plain single-column rewrite you can copy or download as TXT/DOCX. No invented jobs.</p></div>
        </div>
        <div class="section-label">How it works</div>
        <div class="steps">
          <div class="step"><div class="n">01</div><h4>Upload resume</h4><p>PDF or DOCX. Processed in memory, not stored.</p></div>
          <div class="step"><div class="n">02</div><h4>Paste the job</h4><p>Naukri, LinkedIn, or the company JD.</p></div>
          <div class="step"><div class="n">03</div><h4>AI scores fit</h4><p>Gemini returns structured insights.</p></div>
          <div class="step"><div class="n">04</div><h4>Edit and apply</h4><p>Use the gaps and questions the same day.</p></div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_results(result: dict) -> None:
    score = int(result.get("ats_score") or 0)
    summary = result.get("score_summary") or ""
    st.markdown(
        f"""
        <div class="score-wrap">
          <div class="score-label">Fit score · AI estimate, not a real ATS</div>
          <div class="score-num">{score}<span style="font-size:1.4rem;color:#9AA3B2"> / 100</span></div>
          <p style="color:#C9C4B6;margin:10px 0 0 0;max-width:720px;">{summary}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.progress(min(max(score, 0), 100) / 100)

    left, right = st.columns(2)
    with left:
        st.markdown("#### Matched")
        st.markdown(chips(result.get("matched_skills"), "in"), unsafe_allow_html=True)
    with right:
        st.markdown("#### Missing")
        st.markdown(chips(result.get("missing_skills"), "out"), unsafe_allow_html=True)

    ats_resume = (result.get("ats_friendly_resume") or "").strip()
    if ats_resume:
        st.markdown("#### ATS-friendly resume")
        st.caption(
            "Single-column rewrite for this JD. Check every fact before you apply. "
            "Paste into Word or upload the DOCX. Fancy templates often fail ATS parses."
        )
        st.text_area(
            "ATS resume text",
            value=ats_resume,
            height=280,
            label_visibility="collapsed",
        )
        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button(
                "Download .txt",
                data=ats_resume.encode("utf-8"),
                file_name="JDFit_ATS_resume.txt",
                mime="text/plain",
                key="dl_ats_txt",
            )
        with dl2:
            st.download_button(
                "Download .docx",
                data=resume_to_docx(ats_resume),
                file_name="JDFit_ATS_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_ats_docx",
            )

    st.markdown("#### Resume edits")
    for item in result.get("resume_suggestions") or []:
        st.markdown(f'<div class="s-item">{item}</div>', unsafe_allow_html=True)

    st.markdown("#### Interview prep")
    for i, item in enumerate(result.get("interview_questions") or [], 1):
        st.markdown(f'<div class="q-item"><b>{i:02d}</b>{item}</div>', unsafe_allow_html=True)

    roles = result.get("recommended_job_roles") or []
    if roles:
        st.markdown("#### Nearby roles")
        st.markdown(
            "".join(f'<span class="role-pill">{role}</span>' for role in roles),
            unsafe_allow_html=True,
        )


def main():
    st.set_page_config(page_title="JDFit — Resume to interview-ready", page_icon="J", layout="wide")
    inject_css()
    if not gate_password():
        return

    if "analyses" not in st.session_state:
        st.session_state.analyses = 0

    render_hero()

    resume_file = st.file_uploader("Resume (PDF or DOCX)", type=["pdf", "docx"])
    jd_text = st.text_area(
        "Paste the job description",
        height=180,
        placeholder="Paste the full JD from Naukri, LinkedIn, or the company site.",
    )

    go = st.button("Analyze fit — it's free")
    if go:
        if st.session_state.analyses >= MAX_ANALYSES_PER_SESSION:
            st.warning("Session limit reached. Refresh later so the free Gemini quota lasts for other testers.")
        elif not resume_file:
            st.error("Upload a resume.")
        elif not jd_text or len(jd_text.strip()) < 40:
            st.error("Paste a fuller job description.")
        else:
            try:
                with st.spinner("Scoring resume against the JD…"):
                    resume_text = extract_resume(resume_file)
                    result = analyze(resume_text, jd_text.strip())
                st.session_state.analyses += 1
                st.session_state.result = result
            except Exception as exc:
                st.error(str(exc))

    result = st.session_state.get("result")
    if result:
        st.markdown('<div class="section-label">Your toolkit</div>', unsafe_allow_html=True)
        render_results(result)

    render_below_fold()

    st.markdown(
        '<p class="foot">JDFit MVP · resumes processed in memory, not saved · '
        "do not upload documents you cannot share · review the ATS rewrite before you apply</p>",
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
